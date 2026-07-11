"""DOCX extractor."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from ai_os.knowledge.models import ExtractedDocument, ExtractionQuality, Format


class DocxExtractor:
    supported_formats = (Format.DOCX,)

    def extract(self, data: bytes, path: Path | None = None) -> ExtractedDocument:
        document = Document(BytesIO(data))
        blocks: list[str] = []
        title = path.stem if path else "Untitled"

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if "heading 1" in style:
                blocks.append(f"# {text}")
                if title == (path.stem if path else "Untitled"):
                    title = text
            elif "heading 2" in style:
                blocks.append(f"## {text}")
            elif "heading 3" in style:
                blocks.append(f"### {text}")
            else:
                blocks.append(text)

        body = "\n\n".join(blocks).strip()
        return ExtractedDocument(
            title=title,
            body=body,
            extraction_quality=ExtractionQuality.FULL if body else ExtractionQuality.FAILED,
        )
